import os
import json
import streamlit as st
from dotenv import load_dotenv
from docx import Document
from io import BytesIO
from fuzzywuzzy import fuzz
from bs4 import BeautifulSoup
from langchain_community.vectorstores import FAISS
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from docx.shared import RGBColor
import asyncio

# ---------------- Utility Functions ----------------

def clean_html(raw_html):
    """Remove HTML tags & extra whitespace from text."""
    return BeautifulSoup(raw_html, "html.parser").get_text(separator=" ", strip=True)

def shorten_text(text, max_chars=200):
    """Shorten text to max_chars, cutting at sentence boundary if possible."""
    if len(text) <= max_chars:
        return text
    cut = text[:max_chars].rsplit('.', 1)[0]
    return cut.strip() + "..."

# Ensure asyncio loop exists for Google API
try:
    asyncio.get_running_loop()
except RuntimeError:
    asyncio.set_event_loop(asyncio.new_event_loop())

# ---------------- Load Environment & Vectorstore ----------------
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=api_key)
vectorstore = FAISS.load_local("adgm_vectorstore", embeddings, allow_dangerous_deserialization=True)

# ---------------- Checklists ----------------
CHECKLISTS = {
    "Articles of Association": [
        "Company name, type, and registered address clearly stated",
        "Share capital and number of shares specified",
        "Rights attached to each class of shares described",
        "Share transfer restrictions, if any, included",
        "Procedures for shareholder meetings outlined",
        "Voting rights and decision-making process explained",
        "Appointment, removal, and powers of directors detailed",
        "Dividend distribution policy defined",
        "Procedures for amending the Articles specified",
        "Winding up or dissolution clauses included"
    ],
    "Service Agreement": [
        "Names and contact details of both parties included",
        "Detailed description of services to be provided",
        "Payment terms, schedule, and method specified",
        "Start date and duration of the agreement stated",
        "Termination clauses, including notice period, included",
        "Confidentiality obligations outlined",
        "Dispute resolution mechanism stated",
        "Governing law specified",
        "Liability limitations and indemnities included",
        "Force majeure clause included"
    ],
    "Privacy Policy": [
        "Clear statement of data collection practices",
        "Types of personal data collected listed",
        "Purpose for collecting and processing data explained",
        "Legal basis for processing data stated",
        "Data retention period specified",
        "User rights under applicable law described",
        "Information on third-party sharing included",
        "Security measures for data protection outlined",
        "Contact details for privacy inquiries provided",
        "Procedure for policy updates described"
    ]
}

# ---------------- Core Logic ----------------
def detect_document_type(text):
    """Detect document type based on keywords."""
    text_lower = text.lower()
    aoa_keywords = ["articles of association", "shareholder", "board of directors", "company constitution"]
    service_keywords = ["service agreement", "services provided", "payment terms", "termination clause"]
    privacy_keywords = ["privacy policy", "personal data", "data protection", "user rights"]

    def keyword_score(keywords):
        return sum(1 for kw in keywords if kw in text_lower)

    scores = {
        "Articles of Association": keyword_score(aoa_keywords),
        "Service Agreement": keyword_score(service_keywords),
        "Privacy Policy": keyword_score(privacy_keywords)
    }
    detected_type = max(scores, key=scores.get)
    return detected_type if scores[detected_type] > 0 else "Unknown"

def run_checklist(doc_type, text):
    """Match checklist items and return missing ones."""
    issues = []
    if doc_type in CHECKLISTS:
        for item in CHECKLISTS[doc_type]:
            if fuzz.partial_ratio(item.lower(), text.lower()) < 70:
                results = vectorstore.similarity_search(item, k=1)
                if results:
                    raw_citation = clean_html(results[0].page_content)
                    citation_short = shorten_text(raw_citation, 200)
                else:
                    raw_citation = "No citation found"
                    citation_short = raw_citation
                issues.append({
                    "requirement": item,
                    "citation": citation_short,
                    "citation_full": raw_citation
                })
    return issues

def annotate_docx(uploaded_file, issues):
    """Add inline annotations to DOCX for each missing checklist item."""
    doc = Document(uploaded_file)
    for issue in issues:
        found = False
        for para in doc.paragraphs:
            if any(word in para.text.lower() for word in issue["requirement"].lower().split()):
                note = f" [⚠ Missing: {issue['requirement']} — Related Law: {issue['citation']}]"
                para.add_run(note).font.color.rgb = RGBColor(255, 0, 0)
                found = True
                break
        if not found:
            p = doc.add_paragraph()
            p.add_run(f"[⚠ Missing: {issue['requirement']} — Related Law: {issue['citation']}]").font.color.rgb = RGBColor(255, 0, 0)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ---------------- Streamlit UI ----------------
st.set_page_config(page_title="Compliance Checker", layout="wide")
st.title("📄 Document Compliance Checker with ADGM Laws")

uploaded_file = st.file_uploader("Upload a .docx file", type=["docx"])

if uploaded_file:
    doc = Document(uploaded_file)
    full_text = "\n".join([p.text for p in doc.paragraphs])

    doc_type = detect_document_type(full_text)
    st.subheader(f"Detected Document Type: **{doc_type}**")

    issues = run_checklist(doc_type, full_text)

    if issues:
        st.error(f"⚠ {len(issues)} missing or non-compliant items found:")
        for issue in issues:
            st.markdown(f"- **{issue['requirement']}** — _Related Law_: {issue['citation']}")
        
        annotated_file = annotate_docx(uploaded_file, issues)
        st.download_button(
            label="⬇ Download Annotated DOCX",
            data=annotated_file,
            file_name="annotated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        json_output = json.dumps({"doc_type": doc_type, "issues": issues}, indent=2, ensure_ascii=False)
        st.download_button("⬇ Download JSON Report", json_output, "report.json")
    else:
        st.success("✅ Document meets all checklist requirements!")


